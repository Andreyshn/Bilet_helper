from docx import Document
import copy

def ex(doc):
    for elem in doc.element.body:
        if elem.tag.endswith('tbl'):
            return elem
    else:
        print('Неверный формат шаблонного документа!')
        return 1
    
def fill(rdy, theme_count, table, newDoc):
    g = copy.deepcopy(table)
    for i in range(len(rdy)):
        newDoc._body._element._insert_tbl(table)
        del table
        table = copy.deepcopy(g)
        newDoc.add_paragraph('\n')
        newDoc.add_heading(f'Билет № {i+1}', level = 2)
        for j in range(theme_count):
            newDoc.add_paragraph(f'{(j)%theme_count+1}. {rdy[i][j]}')
        newDoc.add_paragraph('\n')
        newDoc.save('ready.docx')

def ins(rdy, theme_count):
    temp = Document('template.docx')
    newDoc = Document()
    table = ex(temp)
    fill(rdy, theme_count, table, newDoc)