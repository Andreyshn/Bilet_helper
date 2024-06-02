from docx.shared import Pt
from docx import Document
import random as rnd
import time as t
from insert import ins

class Question:
    def __init__(self,question,topic,importance):
        self.question = question # вопрос
        self.topic = topic #номер темы
        self.importance = importance #сложность

def main():
    d = [] #список объектов вопрос
    theme = [0.0] #список сложности тем
    try:
        print("Введите количество билетов: ")
        bilet_count = int(input())
        print("Введите количество вопросов в билете: ")
        theme_count = int(input()) #количество вопросов
    except ValueError:
        print("Неправильный ввод")
        print("Пожалуйста, введите верное количество")
        return 1
    
    full_text = getText('voprosy.docx')
    #final = Document('ready.docx')
    theme_in_doc= 1  # Счетчик количества тем, добавленных в словарь
    q_imp = 0 #Сложность вопроса
    count_q = 0 #счетчик вопросов каждой отдельной темы
    for i in range(len(full_text)):
        if full_text[i] == '':  # пустая строка - переход на другую тему
            theme[theme_in_doc-1] /= count_q #среднее значение сложности темы
            theme_in_doc+= 1
            count_q = 0
            theme.append(0)
        else:
            try:
                q_imp = int(full_text[i][-1])
                d.append(Question(full_text[i][:-1],theme_in_doc,q_imp)) #сложность - последняя цифра вопроса
            except ValueError:
                q_imp = 5
                d.append(Question(full_text[i],theme_in_doc,q_imp)) #если сложность не задана - равна 5
            finally:
                theme[theme_in_doc-1] += q_imp
                count_q += 1
    theme[theme_in_doc-1] /= count_q #среднее значение сложности темы
    if theme_count == theme_in_doc:
        rdy = equality(d,theme_in_doc,bilet_count) #равенство тем и вопросов в билете
    elif theme_count > theme_in_doc:
        rdy = more(d,theme_in_doc,bilet_count) #тем больше вопросов
    else:
        rdy = less(d,theme_in_doc,bilet_count) #тем меньше вопросов
    print(ins(rdy, theme_count))
    print(theme)
    t.sleep(10)
    return 0


def less(d,theme_in_doc, N):
    pass

def more(d,theme_in_doc,N):
    pass


def equality(d,theme_in_doc,N):
    rdy = []
    for i in range(1,theme_in_doc+1):
            key = i 
            q = [obj.question for obj in d if obj.topic == i]
            imp = [obj.importance for obj in d if obj.topic == i]
            for _ in range(N):
                key = generateTicket(q,imp,rdy,key,i)
    return rdy

def generateTicket(q,imp,rdy, key,key_copy):
    if key == 1:
        totalProbability = sum(imp)
        rnd_value = rnd.randint(1,totalProbability)
        cumulativeProbability = 0
        for i, prob in enumerate(imp):
            cumulativeProbability += prob
            if rnd_value <= cumulativeProbability:
                rdy.append(q[i])
                return key
    else:
        totalProbability = sum(imp)
        rnd_value = rnd.randint(1,totalProbability)
        cumulativeProbability = 0
        for i, prob in enumerate(imp):
            cumulativeProbability += prob
            if rnd_value <= cumulativeProbability:
                rdy.insert(key-1,q[i])
                return key+key_copy
def getText(filename):
    doc = Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return full_text

def formatDoc(doc):
    doc.styles['Normal'].font.name = 'Times New Roman'  # Название шрифта
    doc.styles['Normal'].font.size = Pt(12)  # Размер шрифта
    
main()